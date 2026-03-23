from __future__ import annotations

import copy
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from models import ParsedTemplate
from utils import SECTION_KEYWORDS, WarningCollector, normalize_text


def _replace_paragraph_text_minimal(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.text = text


def _replace_cell_text_minimal(cell: _Cell, text: str) -> None:
    if not cell.paragraphs:
        cell.text = text
        return
    _replace_paragraph_text_minimal(cell.paragraphs[0], text)
    for p in cell.paragraphs[1:]:
        _replace_paragraph_text_minimal(p, "")


def _row_visual_cells(row_xml: Any) -> list[tuple[Any, int]]:
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tc_tag = f"{{{w_ns}}}tc"
    tcpr_tag = f"{{{w_ns}}}tcPr"
    gs_tag = f"{{{w_ns}}}gridSpan"
    val_attr = f"{{{w_ns}}}val"
    result: list[tuple[Any, int]] = []
    for tc in row_xml.findall(tc_tag):
        span = 1
        tc_pr = tc.find(tcpr_tag)
        if tc_pr is not None:
            gs = tc_pr.find(gs_tag)
            if gs is not None:
                span = int(gs.get(val_attr, "1"))
        result.append((tc, span))
    return result


def _resolve_table_payload(table_type: str, tables_payload: Any) -> list[dict[str, Any]]:
    if isinstance(tables_payload, dict):
        rows = tables_payload.get(table_type) or tables_payload.get("unknown") or []
        return rows if isinstance(rows, list) else []
    if isinstance(tables_payload, list):
        for item in tables_payload:
            if item.get("type") == table_type and isinstance(item.get("rows"), list):
                return item["rows"]
    return []


class DocxInjectionEngine:
    def __init__(self, template_path: str, parsed: ParsedTemplate, warnings: WarningCollector) -> None:
        self.document = Document(template_path)
        self.parsed = parsed
        self.warnings = warnings

    def inject_sections(self, content: dict[str, str], max_chars_per_section: int = 1400) -> None:
        paragraphs = self.document.paragraphs
        for section, text in content.items():
            target_text = text.strip()
            if len(target_text) > max_chars_per_section:
                target_text = target_text[:max_chars_per_section]
                self.warnings.add("section", "페이지 초과 가능성: 장문 텍스트가 잘렸습니다.", section=section)

            injected = False
            placeholder = f"{{{{{section}}}}}"
            for p in paragraphs:
                if placeholder in p.text:
                    _replace_paragraph_text_minimal(p, p.text.replace(placeholder, target_text))
                    injected = True
                    break

            if injected:
                continue

            keywords = [section] + SECTION_KEYWORDS.get(section, [])
            for idx, p in enumerate(paragraphs):
                if any(k.lower() in p.text.lower() for k in keywords):
                    if idx + 1 < len(paragraphs):
                        _replace_paragraph_text_minimal(paragraphs[idx + 1], target_text)
                        injected = True
                        break
            if not injected:
                self.warnings.add("section", "섹션 주입 위치를 찾지 못했습니다.", section=section)

    def inject_tables(self, tables_payload: Any) -> None:
        for diag in self.parsed.table_diagnostics:
            if diag.table_index >= len(self.document.tables):
                continue
            table = self.document.tables[diag.table_index]
            rows_payload = _resolve_table_payload(diag.table_type_candidate, tables_payload)
            if not rows_payload:
                continue

            if diag.table_risk_level == "manual":
                self.warnings.add(
                    "table",
                    "병합셀(vMerge) 구조로 자동주입을 건너뜁니다. 수동 확인 필요",
                    table_index=diag.table_index,
                    reason=diag.reason,
                )
                continue

            self._inject_table_rows(table, rows_payload, diag)

    def _inject_table_rows(self, table: Table, rows_payload: list[dict[str, Any]], diag: Any) -> None:
        if len(table.rows) < 2:
            self.warnings.add("table", "샘플 데이터행이 없어 표 자동주입 불가", table_index=diag.table_index)
            return

        sample_row_xml = table.rows[1]._tr
        tbl_xml = table._tbl

        while len(table.rows) > 2:
            tbl_xml.remove(table.rows[-1]._tr)

        for idx, row_item in enumerate(rows_payload):
            target_row = sample_row_xml if idx == 0 else copy.deepcopy(sample_row_xml)
            if idx > 0:
                tbl_xml.append(target_row)

            payload_values = list(row_item.values()) if isinstance(row_item, dict) else list(row_item)
            row_cells = _row_visual_cells(target_row)
            val_idx = 0
            for tc, span in row_cells:
                text = payload_values[val_idx] if val_idx < len(payload_values) else ""
                cell_obj = _Cell(tc, table)
                _replace_cell_text_minimal(cell_obj, str(text))
                val_idx += max(1, span)

    def inject_images(self, images_manifest: list[dict[str, Any]], image_root: str | None = None) -> None:
        paragraphs = self.document.paragraphs
        for idx, item in enumerate(images_manifest):
            section = str(item.get("section", "")).strip().lower()
            caption = str(item.get("caption", "")).strip()
            image_path = Path(image_root, item["filename"]) if image_root else Path(item["filename"])
            if not image_path.exists():
                self.warnings.add("image", "이미지 파일이 존재하지 않습니다.", image=str(image_path))
                continue

            anchor = None
            for p in paragraphs:
                text_l = p.text.lower()
                if section and section in text_l:
                    anchor = p
                    break
                if caption and caption.lower() in text_l:
                    anchor = p
                    break
                if "[image]" in text_l or "[이미지]" in text_l:
                    anchor = p
                    break

            if anchor is None:
                for i, p in enumerate(paragraphs):
                    if section and any(k in p.text.lower() for k in [section, *SECTION_KEYWORDS.get(section, [])]):
                        anchor = paragraphs[min(i + 1, len(paragraphs) - 1)]
                        self.warnings.add(
                            "image",
                            "자동 삽입 위치 미확정, 섹션 끝 fallback 사용",
                            section=section,
                            image_index=idx,
                        )
                        break

            if anchor is None:
                anchor = self.document.add_paragraph()
                self.warnings.add("image", "앵커 없음: 문서 끝에 fallback 삽입", image_index=idx)

            run = anchor.add_run()
            try:
                run.add_picture(str(image_path))
            except Exception as exc:
                self.warnings.add("image", "이미지 삽입 실패", error=str(exc), image=str(image_path))
                continue

            if caption:
                cap_para = anchor.insert_paragraph_before(caption)
                _replace_paragraph_text_minimal(cap_para, caption)

    def save(self, output_path: str) -> None:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        self.document.save(output_path)
