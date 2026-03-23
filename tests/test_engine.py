from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import RGBColor

from injector import DocxInjectionEngine
from models import ParsedTemplate
from parser import analyze_template
from postprocess import cleanup_guidance_phrases
from utils import WarningCollector


def _save(doc: Document, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    return path


def test_case_a_simple_repeated_table_injection(tmp_path: Path) -> None:
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    t.cell(0, 0).text = "기간"
    t.cell(0, 1).text = "업무"
    t.cell(0, 2).text = "담당"
    t.cell(1, 0).text = "샘플"
    t.cell(1, 1).text = "샘플"
    t.cell(1, 2).text = "샘플"
    path = _save(doc, tmp_path / "a.docx")

    parsed = analyze_template(str(path))
    wc = WarningCollector()
    engine = DocxInjectionEngine(str(path), parsed, wc)
    engine.inject_tables({"team": [{"기간": "1Q", "업무": "개발", "담당": "A"}]})

    assert len(engine.document.tables[0].rows) == 2
    assert "1Q" in engine.document.tables[0].cell(1, 0).text


def test_case_b_gridspan_table(tmp_path: Path) -> None:
    doc = Document()
    t = doc.add_table(rows=3, cols=4)
    t.cell(0, 0).merge(t.cell(0, 1))
    t.cell(0, 0).text = "구분"
    t.cell(0, 2).text = "담당업무"
    t.cell(0, 3).text = "보유역량"
    _save(doc, tmp_path / "b.docx")

    parsed = analyze_template(str(tmp_path / "b.docx"))
    assert parsed.table_diagnostics[0].has_gridspan is True
    assert parsed.table_diagnostics[0].table_risk_level in {"caution", "manual"}


def test_case_c_vmerge_detect_manual(tmp_path: Path) -> None:
    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "조직"
    t.cell(1, 0).merge(t.cell(2, 0))
    _save(doc, tmp_path / "c.docx")

    parsed = analyze_template(str(tmp_path / "c.docx"))
    assert parsed.table_diagnostics[0].has_vmerge is True
    assert parsed.table_diagnostics[0].table_risk_level == "manual"


def test_case_d_guidance_cleanup(tmp_path: Path) -> None:
    doc = Document()
    p = doc.add_paragraph("예시: 이 문구는 삭제 후 작성")
    p.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    p2 = doc.add_paragraph("실제 본문")
    path = _save(doc, tmp_path / "d.docx")

    parsed = ParsedTemplate(template_path=str(path))
    wc = WarningCollector()
    engine = DocxInjectionEngine(str(path), parsed, wc)
    cleanup_guidance_phrases(engine.document, wc)
    texts = [p.text for p in engine.document.paragraphs]
    assert "실제 본문" in texts


def test_case_e_image_fallback_warning(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("1. problem")
    path = _save(doc, tmp_path / "e.docx")

    import pytest
    Image = pytest.importorskip("PIL.Image")

    img_path = tmp_path / "img.png"
    Image.new("RGB", (20, 20), color=(255, 0, 0)).save(img_path)

    parsed = analyze_template(str(path))
    wc = WarningCollector()
    engine = DocxInjectionEngine(str(path), parsed, wc)
    engine.inject_images([{"filename": "img.png", "section": "team", "caption": "팀 사진"}], image_root=str(tmp_path))
    assert any(w["category"] == "image" for w in wc.to_list())


def test_case_f_long_text_warning(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("team")
    doc.add_paragraph("placeholder")
    path = _save(doc, tmp_path / "f.docx")

    parsed = analyze_template(str(path))
    wc = WarningCollector()
    engine = DocxInjectionEngine(str(path), parsed, wc)
    engine.inject_sections({"team": "가" * 2000}, max_chars_per_section=500)
    assert any("페이지 초과" in w["message"] for w in wc.to_list())
