from __future__ import annotations

import re

from docx.document import Document as DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.shared import RGBColor

from utils import WarningCollector

GUIDE_PATTERNS = ["삭제 후", "예시", "작성요령", "참고", "기재", "해당 시", "안내", "※"]
TOC_START = ["목차", "차례", "contents"]
TAIL_ANCHORS = ["사업비 집행 비목 참고", "삭제 후 제출", "참고사항"]


def _is_blue_run(run) -> bool:
    rgb = run.font.color.rgb
    if rgb is None:
        return False
    r, g, b = rgb[0], rgb[1], rgb[2]
    return b > 120 and b >= r and b >= g


def _remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def cleanup_guidance_phrases(doc: DocxDocument, warnings: WarningCollector) -> None:
    for p in list(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        has_pattern = any(token in text for token in GUIDE_PATTERNS)
        blue_count = 0
        for run in p.runs:
            if _is_blue_run(run):
                run.font.color.rgb = RGBColor(0, 0, 0)
                blue_count += 1

        if has_pattern and (blue_count > 0 or len(text) < 80):
            if re.search(r"(예시|작성요령|삭제 후|참고)", text):
                _remove_paragraph(p)
                warnings.add("postprocess", "안내문구 문단 제거", text=text[:60])


def remove_toc_block(doc: DocxDocument, warnings: WarningCollector) -> None:
    paras = list(doc.paragraphs)
    start_idx = None
    for i, p in enumerate(paras):
        if p.text.strip().lower() in TOC_START or any(k in p.text.strip().lower() for k in TOC_START):
            start_idx = i
            break
    if start_idx is None:
        return

    end_idx = None
    section_re = re.compile(r"^\s*\d+(?:[-\.)]\d+)*")
    for i in range(start_idx + 1, len(paras)):
        if section_re.match(paras[i].text.strip()) and len(paras[i].text.strip()) > 2:
            end_idx = i
            break
    if end_idx is None:
        warnings.add("postprocess", "목차 블록 제거 경계 탐지 실패")
        return

    for i in range(start_idx, end_idx):
        _remove_paragraph(paras[i])
    warnings.add("postprocess", "목차 블록 제거 수행", start=start_idx, end=end_idx)


def remove_tail_reference_block(doc: DocxDocument, warnings: WarningCollector) -> None:
    paras = list(doc.paragraphs)
    anchor_idx = None
    for i, p in enumerate(paras):
        if any(a in p.text for a in TAIL_ANCHORS):
            anchor_idx = i
    if anchor_idx is None:
        return

    for p in paras[anchor_idx:]:
        _remove_paragraph(p)
    warnings.add("postprocess", "끝부분 참고 블록 제거", start=anchor_idx)
